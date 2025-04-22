from docx import Document
import json

def create_docx(questions, filename):
    # 创建 Word 文档对象
    doc = Document()

    # 添加标题
    doc.add_heading('医学问题与解析', level=1)

    # 遍历每个问题
    for i, question in enumerate(questions, 1):
        try:
            # 添加问题编号和题干
            doc.add_heading(f'问题 {i}', level=2)
            doc.add_paragraph(question['topic'])

            # 检查是否为子问题列表
            if 'questions' in question:
                # 对于包含子问题的大题目，逐个处理子问题
                for j, sub_question in enumerate(question['questions'], 1):
                    doc.add_heading(f'子问题 {i}.{j}', level=3)
                    doc.add_paragraph(sub_question['question'])

                    # 添加选项
                    options = sub_question.get('options', {})
                    for key, value in options.items():
                        doc.add_paragraph(f"{key}. {value}", style='List Bullet')

                    # 添加答案
                    doc.add_heading('答案', level=4)
                    doc.add_paragraph(f"正确答案: {sub_question.get('answer', '无答案')}")

                    # 添加解析
                    doc.add_heading('解析', level=4)
                    doc.add_paragraph(sub_question.get('parse', '无解析'))
            else:
                # 对于普通问题，直接处理
                options = question.get('options', {})
                for key, value in options.items():
                    doc.add_paragraph(f"{key}. {value}", style='List Bullet')

                # 添加答案
                doc.add_heading('答案', level=3)
                doc.add_paragraph(f"正确答案: {question.get('answer', '无答案')}")

                # 添加解析
                doc.add_heading('解析', level=3)
                doc.add_paragraph(question.get('parse', '无解析'))

        except KeyError as e:
            print(f"错误: 问题 {i} 缺少必要字段: {e}")
        except Exception as e:
            print(f"问题 {i} 处理失败: {e}")

    # 保存文档
    doc.save(filename)


with open('../save/questionGeneration.json', 'r', encoding='utf-8') as f:
    questions = json.load(f)

# 生成 Word 文档
output_filename = "../save/医学问题与解析.docx"
create_docx(questions, output_filename)
print(f"文档已保存为 {output_filename}")
